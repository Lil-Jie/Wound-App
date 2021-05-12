from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from db.sqlite3_connect import insert_data, select_data, create_table
import time


def word_file(file):
    sql = 'SELECT * FROM report WHERE id=1'
    res = select_data(sql)
    name = res[0][1]
    gender = res[0][2]
    age = res[0][3]
    height = res[0][4]
    weight = res[0][5]
    area = res[0][6]
    perimeter = res[0][7]

    doc = Document()

    doc.add_heading("Condition Report", 0)

    p = doc.add_paragraph(' This report is about your ')
    p.add_run('wound detection').bold = True
    p.add_run(' and ')
    p.add_run('issue analysis.').bold = True

    doc.add_heading('Personal Information', level=1)
    doc.add_paragraph("Name:  "+name, style='List Bullet')
    doc.add_paragraph("Gender:  "+gender, style='List Bullet')
    doc.add_paragraph("Age:  "+age, style='List Bullet')
    doc.add_paragraph("Height:  "+height+" cm", style='List Bullet')
    doc.add_paragraph("Weight:  "+weight+" kg", style='List Bullet')

    p2 = doc.add_paragraph('Wound Classification：')
    p2.add_run('Burns').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_heading('Diagnostic Results', 1)
    table = doc.add_table(rows=2, cols=2, style='Table Grid')
    first_cells = table.rows[0].cells
    first_cells[0].text = 'Original Image'
    first_cells[1].text = 'Segmented Image'
    run = doc.tables[0].cell(1, 0).paragraphs[0].add_run()
    run.add_picture("%s" % file, width=Inches(1.5))
    run = doc.tables[0].cell(1, 1).paragraphs[0].add_run()
    run.add_picture('appimage/wound_image/output_wound/result.png', width=Inches(1.5))
    doc.add_paragraph("Wound area: " + area, style='List Bullet')
    doc.add_paragraph("Wound perimeter: " + perimeter, style='List Bullet')
    doc.add_paragraph("Degree of wound color:", style='List Bullet')

    doc.add_heading('Diagnostic Suggestion', 1)
    doc.add_paragraph("Immediately rinse with cold water can effectively relieve pain and reduce edema.", style='List Number')
    doc.add_paragraph("Take off the clothes carefully in cold water. Use alcohol-sterilized scissors to cut along the skin around the skin. Don't pull it hard.", style='List Number')
    doc.add_paragraph("Clean burn and scald wounds in time, keep the wound clean and hygienic to prevent infection.", style='List Number')
    doc.add_paragraph("When blisters appear on the skin, don't break the blisters. After covering the wound with clean gauze, send it to the hospital for treatment immediately.", style='List Number')

    doc.add_paragraph('This report is for reference only and does not provide any other certification. Please ask professional doctor for details.', style='Intense Quote')
    doc.add_paragraph('\t\t\t\t\t\t\t\t\t Date:'+time.strftime("%d/%m/%Y"))
    doc.save("Condition Report.docx")
